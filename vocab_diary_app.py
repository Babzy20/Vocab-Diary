import streamlit as st
import pandas as pd
import requests
from docx import Document
from io import BytesIO

# Enable wide layout
st.set_page_config(layout="wide")
st.title("Vocabulary Diary")

# Initialize session state for word history
if "word_history" not in st.session_state:
    st.session_state.word_history = []

def fetch_word_details(word):
    url = f"https://api.dictionaryapi.dev/api/v2/entries/en/{word}"
    response = requests.get(url)
    dictionary_link = f"https://www.merriam-webster.com/dictionary/{word}"

    if response.status_code != 200:
        return {
            "Word": word,
            "Definition": f'Definition not found. <a href="{dictionary_link}" target="_blank">Search on Merriam-Webster</a>',
            "Example Sentence": f'No example available. <a href="{dictionary_link}" target="_blank">Search on Merriam-Webster</a>',
            "IPA": f'IPA not found. <a href="{dictionary_link}" target="_blank">Search on Merriam-Webster</a>',
            "Audio URL": ""
        }

    try:
        data = response.json()[0]
        definition = data["meanings"][0]["definitions"][0].get("definition", f'Definition not found. <a href="{dictionary_link}" target="_blank">Search on Merriam-Webster</a>')
        example = data["meanings"][0]["definitions"][0].get("example", f'No example available. <a href="{dictionary_link}" target="_blank">Search on Merriam-Webster</a>')
        ipa = data.get("phonetic", "")
        if not ipa:
            ipa = next((p.get("text", "") for p in data.get("phonetics", []) if "text" in p), "")
            if not ipa:
                ipa = f'IPA not found. <a href="{dictionary_link}" target="_blank">Search on Merriam-Webster</a>'
        audio_url = next((p.get("audio", "") for p in data.get("phonetics", []) if "audio" in p and p["audio"]), "")
    except Exception:
        return {
            "Word": word,
            "Definition": f'Definition not found. <a href="{dictionary_link}" target="_blank">Search on Merriam-Webster</a>',
            "Example Sentence": f'No example available. <a href="{dictionary_link}" target="_blank">Search on Merriam-Webster</a>',
            "IPA": f'IPA not found. <a href="{dictionary_link}" target="_blank">Search on Merriam-Webster</a>',
            "Audio URL": ""
        }

    return {
        "Word": word,
        "Definition": definition,
        "Example Sentence": example,
        "IPA": ipa,
        "Audio URL": audio_url
    }

def create_word_document(df, include_definition, include_example, include_ipa, include_audio):
    doc = Document()
    doc.add_heading('Vocabulary Diary', 0)

    for index, row in df.iterrows():
        doc.add_heading(row['Word'], level=1)
        if include_definition and "Definition" in row:
            doc.add_paragraph(f"Definition: {row['Definition']}")
        if include_example and "Example Sentence" in row:
            doc.add_paragraph(f"Example Sentence: {row['Example Sentence']}")
        if include_ipa and "IPA" in row:
            doc.add_paragraph(f"IPA: {row['IPA']}")
        if include_audio and "Audio URL" in row and row['Audio URL']:
            doc.add_paragraph(f"Audio URL: {row['Audio URL']}")
        if "Personal Definition" in row and row["Personal Definition"]:
            doc.add_paragraph(f"My Personal Definition: {row['Personal Definition']}")

    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

@st.cache_data
def fetch_word_details_cached(word):
    return fetch_word_details(word)

# Sidebar for toggle options
st.sidebar.title("Settings")
include_definition = st.sidebar.checkbox("Include Definition", value=True)
include_example = st.sidebar.checkbox("Include Example Sentence", value=True)
include_ipa = st.sidebar.checkbox("Include IPA", value=True)
include_audio = st.sidebar.checkbox("Include Audio URL", value=True)

# Input area
words_input = st.text_area("Enter words (comma, space, or new line separated):")

if st.button("Fetch Word Details"):
    words = [word.strip() for word in words_input.replace(',', ' ').split()]
    word_details = []

    for word in words:
        details = fetch_word_details_cached(word)
        if details:
            word_details.append(details)
            st.session_state.word_history.append(details)

    if word_details:
        df = pd.DataFrame(word_details)
        
        # Filter columns based on sidebar settings
        columns_to_include = ["Word"]
        if include_definition:
            columns_to_include.append("Definition")
        if include_example:
            columns_to_include.append("Example Sentence")
        if include_ipa:
            columns_to_include.append("IPA")
        if include_audio:
            columns_to_include.append("Audio URL")
        
        df = df[columns_to_include]

        if include_audio:
            df["Audio URL"] = df["Audio URL"].apply(
                lambda url: f'<a href="{url}" target="_blank">🔊 Listen</a>' if url else "No audio"
            )

        st.markdown(df.to_html(escape=False, index=False), unsafe_allow_html=True)
        st.success("✅ Words fetched successfully!")

# Editable word history
if "word_history" in st.session_state and st.session_state.word_history:
    st.subheader("📜 Word History (Editable)")

    for i, word_data in enumerate(st.session_state.word_history):
        st.markdown(f"### {word_data['Word']}")
        if include_definition:
            new_def = st.text_area(f"Definition for {word_data['Word']}", word_data['Definition'], key=f"def_{i}")
            st.session_state.word_history[i]['Definition'] = new_def
        if include_example:
            new_ex = st.text_area(f"Example for {word_data['Word']}", word_data['Example Sentence'], key=f"ex_{i}")
            st.session_state.word_history[i]['Example Sentence'] = new_ex
        if include_ipa:
            new_ipa = st.text_area(f"IPA for {word_data['Word']}", word_data['IPA'], key=f"ipa_{i}")
            st.session_state.word_history[i]['IPA'] = new_ipa

        # New: Personal Definition section
        with st.expander("My Personal Definition", expanded=False):
            personal_def_key = f"personal_def_{i}"
            personal_def = st.text_area("Write your own definition:", word_data.get("Personal Definition", ""), key=personal_def_key)
            st.session_state.word_history[i]["Personal Definition"] = personal_def

    # Prepare DataFrame for download
    df_history = pd.DataFrame(st.session_state.word_history)

    # Filter columns based on sidebar settings
    columns_to_include = ["Word"]
    if include_definition:
        columns_to_include.append("Definition")
    if include_example:
        columns_to_include.append("Example Sentence")
    if include_ipa:
        columns_to_include.append("IPA")
    if include_audio:
        columns_to_include.append("Audio URL")
    if any("Personal Definition" in word for word in st.session_state.word_history):
        columns_to_include.append("Personal Definition")

    df_history = df_history[columns_to_include]

    if include_audio:
        df_history["Audio URL"] = df_history["Audio URL"].apply(
            lambda url: f'<a href="{url}" target="_blank">🔊 Listen</a>' if url else "No audio"
        )

    st.markdown(df_history.to_html(escape=False, index=False), unsafe_allow_html=True)

    st.download_button(
        label="Download History as CSV",
        data=df_history.to_csv(index=False),
        file_name='vocabulary_history.csv',
        mime='text/csv'
    )

    word_doc = create_word_document(df_history, include_definition, include_example, include_ipa, include_audio)
    st.download_button(
        label="Download History as Word Document",
        data=word_doc,
        file_name='vocabulary_history.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    if st.button("Clear Word History"):
        st.session_state.word_history = []
        st.success("🧹 Word history cleared.")

